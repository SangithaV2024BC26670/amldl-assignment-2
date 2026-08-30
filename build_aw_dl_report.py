"""
Builds the deep-learning report (.docx) from the artefacts written by
aw_deep_learning.py, cross-referenced against the benchmark gradient-boosting
results in outputs/tables.

The document is written to stand alone: the gradient-boosting model is referred
to as the benchmark throughout rather than by project phase.
"""

import json
import os

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DL_TAB = os.path.join("outputs_dl", "tables")
DL_FIG = os.path.join("outputs_dl", "figures")
BM_TAB = os.path.join("outputs", "tables")
OUT = "AMLDL_DeepLearning_Report.docx"
ACCENT = RGBColor(0x1F, 0x3B, 0x73)

D = json.load(open(os.path.join(DL_TAB, "dl_t06_run_summary.json")))
BM = json.load(open(os.path.join(BM_TAB, "aw_t15_run_summary.json")))
compare = pd.read_csv(os.path.join(DL_TAB, "dl_t02_benchmark_vs_network.csv"))
hist = pd.read_csv(os.path.join(DL_TAB, "dl_t01_training_history.csv"))
corr = pd.read_csv(os.path.join(DL_TAB, "dl_t04_embedding_vs_buyrate_correlation.csv"))
neigh = pd.read_csv(os.path.join(DL_TAB, "dl_t05_embedding_neighbours.csv"))

# Present neutral model labels regardless of how the source table was written
# Defensive: strip any legacy phase suffix if an older table is present
compare["Model"] = compare["Model"].str.replace(r"\s*\(Phase [23]\)", "",
                                                regex=True)
compare = compare.rename(columns={"Delta_vs_Phase2": "Delta_ROC_AUC"})

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table_from_df(df, widths=None, caption=None):
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(df.columns):
        cell = t.rows[0].cells[i]
        cell.text = str(c)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if pd.isna(v) else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    if caption:
        cap = doc.add_paragraph()
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(8.5)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def figure(fname, caption, width=6.3):
    path = os.path.join(DL_FIG, fname)
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(8.5)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


gb_row = compare.iloc[0]
nn_row = compare.iloc[1]
BM_AUC = D["benchmark_test_ROC_AUC"]
NN_AUC = D["network_test_ROC_AUC"]
DELTA = D["delta_ROC_AUC"]

# ==========================================================================
t = doc.add_heading("Applied Machine Learning and Deep Learning (MBA ZG582)", 0)
for r in t.runs:
    r.font.color.rgb = ACCENT
p = doc.add_paragraph()
r = p.add_run("Deep Learning Phase")
r.bold = True
r.font.size = Pt(13)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
r = p.add_run("Entity-Embedding Neural Network for Bike-Buyer Propensity")
r.font.size = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
r = p.add_run("Adventure Works Cycles | Microsoft AdventureWorks customer data")
r.italic = True
r.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_table(rows=0, cols=2)
meta.style = "Light List Accent 1"
for k, v in [
    ("Framework", D["framework"] + "  (device: " + D["device"] + ")"),
    ("Architecture", D["architecture"]),
    ("Trainable parameters", format(D["trainable_parameters"], ",")),
    ("Embedded features", str(D["embedded_features"]) + " categorical"),
    ("Continuous features", str(D["continuous_features"])),
    ("Data split", str(D["train_rows"]) + " train / " + str(D["val_rows"])
     + " validation / " + str(D["test_rows"]) + " test"),
    ("Training", str(D["epochs_run"]) + " epochs run, best at epoch "
     + str(D["best_epoch"]) + " (early stopping)"),
    ("Neural network ROC-AUC", format(NN_AUC, ".4f")),
    ("Benchmark model", D["benchmark_model"] + " at " + format(BM_AUC, ".4f")),
]:
    row = meta.add_row().cells
    row[0].text = k
    row[1].text = v
    for pp in row[0].paragraphs:
        for rr in pp.runs:
            rr.bold = True
    for c in row:
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(9.5)

# ==========================================================================
h("1. Objective", 1)
para(
    "The benchmark for this project is a tuned histogram gradient-boosting "
    "classifier, which reaches ROC-AUC " + format(BM_AUC, ".4f") + " on a held-out "
    "20 per cent of customers. This report asks two questions of that result.")
bullets([
    "Can a neural network with learned entity embeddings improve on the boosted "
    "trees, using the same data and the same split?",
    "Independently of the score, does the embedding representation reveal "
    "structure in the categorical variables that a tree ensemble cannot express?",
])
para(
    "The second question matters even if the answer to the first is no. A tree "
    "ensemble can tell you that occupation is important; it cannot tell you which "
    "occupations behave alike. An embedding can.")

h("1.1 Why Entity Embeddings", 2)
para(
    "One-hot encoding places every level of a categorical variable at an equal "
    "distance from every other. Under that encoding the model has no way to "
    "represent the idea that two occupations attract similar customers - each "
    "level is simply a separate, unrelated indicator column. An embedding layer "
    "instead assigns each level a short dense vector, learned by gradient descent "
    "alongside the rest of the network, so that levels which behave similarly for "
    "the prediction task end up close together in the embedding space.")
para(
    "Two practical consequences follow. The representation is more compact than "
    "one-hot for high-cardinality fields, and - more useful here - the learned "
    "vectors can be extracted and plotted, so the structure the network "
    "discovered is open to inspection rather than taken on trust.")

# ==========================================================================
h("2. Architecture and Training", 1)
para(
    "Nine categorical fields are embedded and six continuous fields are passed "
    "through batch normalisation, then concatenated and fed to a two-layer "
    "feed-forward network. Embedding width follows the standard heuristic "
    "min(50, (cardinality + 1) / 2), which gives dimensions of 1 to 3 for the "
    "low-cardinality fields in this dataset.")
table_from_df(pd.DataFrame([
    {"Component": "Embedding layers", "Detail": "9 categorical fields, dims 1-3, concatenated (27 dims total with continuous)"},
    {"Component": "Continuous path", "Detail": "6 features, BatchNorm1d"},
    {"Component": "Hidden layer 1", "Detail": "Linear 128 -> BatchNorm -> ReLU -> Dropout 0.35"},
    {"Component": "Hidden layer 2", "Detail": "Linear 64 -> BatchNorm -> ReLU -> Dropout 0.20"},
    {"Component": "Output", "Detail": "Linear 1, BCEWithLogitsLoss"},
    {"Component": "Class imbalance", "Detail": "pos_weight in the loss - the network equivalent of class_weight='balanced'"},
    {"Component": "Optimiser", "Detail": "AdamW, lr 2e-3, weight decay 1e-4, ReduceLROnPlateau on validation AUC"},
    {"Component": "Early stopping", "Detail": "Patience 15 epochs on validation ROC-AUC"},
    {"Component": "Parameters", "Detail": format(D["trainable_parameters"], ",") + " trainable"},
]), widths=[1.6, 4.7])
para(
    "The count fields - children at home, total children, cars owned - are "
    "embedded rather than passed through as integers. This lets the network learn "
    "its own spacing between, say, two children and three, instead of being forced "
    "into the assumption that the effect is linear in the count.")

h("2.1 Guarding Against Divergence from the Benchmark", 2)
para(
    "For the comparison to be meaningful, the network must see exactly the data "
    "the boosted trees saw. The preparation code mirrors the benchmark pipeline "
    "step for step - the same deduplication rule, the same removal of identifier "
    "and personally identifying columns, the same derived features, the same "
    "random seed and the same stratified 80/20 split - and asserts that the "
    "resulting row count matches 16,404, so any silent drift fails loudly rather "
    "than quietly biasing the result. A validation slice is carved from the "
    "training portion only, so the test set is never seen during training or model "
    "selection.")

figure("dl_fig01_learning_curves.png",
       "Figure 1: Learning curves. Training and validation track closely, and "
       "early stopping selects epoch " + str(D["best_epoch"]) + ".")
para(
    "The curves show no meaningful over-fitting: validation loss flattens rather "
    "than turning upward, and the gap to training loss stays narrow throughout. "
    "The regularisation - dropout, weight decay and batch normalisation - is doing "
    "its job on a dataset of this size.")

# ==========================================================================
h("3. Results", 1)
table_from_df(compare, caption="Table 1: Head-to-head on the identical hold-out "
                               "split.")
figure("dl_fig02_benchmark_vs_network.png",
       "Figure 2: ROC overlay, network confusion matrix, and metric comparison.")

h("3.1 The Headline Finding", 2)
para("The neural network does not beat the gradient-boosting model on ranking "
     "quality.", bold=True)
para(
    "ROC-AUC falls from " + format(BM_AUC, ".4f") + " to " + format(NN_AUC, ".4f")
    + ", a difference of " + format(DELTA, "+.4f") + ". Since ROC-AUC is the "
    "metric the campaign actually depends on - it scores the quality of the "
    "customer ranking - the gradient-boosting model remains the recommended "
    "production choice.")
para(
    "This is a well-documented result rather than a failure of implementation. On "
    "tabular data of this size and shape, gradient-boosted decision trees are a "
    "strong and difficult baseline, and neural networks typically need either "
    "substantially more data, higher-cardinality categorical fields where "
    "embeddings can compress meaningfully, or genuine sequential or interaction "
    "structure before they pull ahead. This dataset has 16,404 rows and "
    "categorical fields with at most six levels, which is close to the least "
    "favourable case for the embedding approach.")

h("3.2 Where the Network Is Better", 2)
para(
    "The comparison is not one-sided, and reporting only the AUC would hide "
    "something real. At the default 0.50 cut-off the network is markedly better on "
    "the metrics that reward finding buyers.")
table_from_df(pd.DataFrame([
    {"Metric": "Recall", "Gradient boosting": format(gb_row["Recall"], ".4f"),
     "Embedding network": format(nn_row["Recall"], ".4f"), "Better": "Network"},
    {"Metric": "F1", "Gradient boosting": format(gb_row["F1"], ".4f"),
     "Embedding network": format(nn_row["F1"], ".4f"), "Better": "Network"},
    {"Metric": "Balanced accuracy", "Gradient boosting": format(gb_row["BalancedAcc"], ".4f"),
     "Embedding network": format(nn_row["BalancedAcc"], ".4f"), "Better": "Network"},
    {"Metric": "Precision", "Gradient boosting": format(gb_row["Precision"], ".4f"),
     "Embedding network": format(nn_row["Precision"], ".4f"), "Better": "Boosting"},
    {"Metric": "Accuracy", "Gradient boosting": format(gb_row["Accuracy"], ".4f"),
     "Embedding network": format(nn_row["Accuracy"], ".4f"), "Better": "Boosting"},
    {"Metric": "ROC-AUC", "Gradient boosting": format(gb_row["ROC_AUC"], ".4f"),
     "Embedding network": format(nn_row["ROC_AUC"], ".4f"), "Better": "Boosting"},
]), widths=[1.5, 1.6, 1.6, 1.0],
    caption="Table 2: Metric-by-metric comparison at the default 0.50 threshold.")
para(
    "The reason is the pos_weight term in the loss, which penalises missed buyers "
    "more heavily and pushes the network to a different natural operating point - "
    "it catches " + format(nn_row["Recall"] * 100, ".1f") + " per cent of buyers "
    "against " + format(gb_row["Recall"] * 100, ".1f") + " per cent, at the cost "
    "of lower precision. That is an artefact of where the threshold sits, not "
    "evidence of a better model: threshold analysis on the benchmark showed that "
    "moving its cut-off to " + format(BM["f1_optimal_threshold"], ".2f")
    + " raises its F1 to 0.6841, which closes most of the gap.")
para(
    "The honest summary is that the two models are close in overall discriminative "
    "power, and the gradient-boosting model retains a small but consistent edge on "
    "the ranking metric that matters commercially.", italic=True)

# ==========================================================================
h("4. What the Embeddings Learned", 1)
para(
    "This is the part of the exercise that produces something the tree ensemble "
    "could not. Each categorical level now has a learned coordinate, and those "
    "coordinates can be plotted and read.")
figure("dl_fig03_learned_embeddings.png",
       "Figure 3: Learned entity embeddings, projected to two dimensions via PCA. "
       "Colour shows the actual purchase rate for each level.")

h("4.1 The Network Organised Categories by Propensity", 2)
para(
    "Nothing in the training procedure instructed the network to arrange each "
    "categorical field along a purchase-propensity axis. It did so anyway. "
    "Correlating the first embedding axis against the observed purchase rate for "
    "each level gives:")
table_from_df(corr, widths=[1.8, 0.9, 1.9, 1.0],
              caption="Table 3: Correlation between embedding axis 1 and the "
                      "observed purchase rate per level.")
para(
    "Education shows the strongest alignment at "
    + format(corr.iloc[0]["Corr_axis1_vs_buyrate"], "+.3f")
    + ", followed by occupation. Country is the weakest, which is consistent with "
    "the earlier finding that geography carries less signal than household "
    "attributes - the network had little propensity structure to encode there, so "
    "it encoded little.")

h("4.2 Which Levels the Network Considers Alike", 2)
table_from_df(neigh, widths=[1.4, 2.4, 2.5],
              caption="Table 4: Closest and most distant level pairs in each "
                      "embedding space.")
para(
    "The education result is the most striking: the closest pair is Partial "
    "College and Bachelors, and the most distant is High School and Graduate "
    "Degree. The network recovered the ordinal structure of the variable from the "
    "data alone, without being told the levels have an order.")
para(
    "One methodological caveat should be stated. The nearest-neighbour distances "
    "in Table 4 are computed in the full embedding space, while Figure 3 shows a "
    "two-dimensional PCA projection of it. Points that appear close in the plot "
    "are not always the closest pair in the underlying space, and the table is the "
    "authoritative one.", italic=True)

h("4.3 Marketing Implications", 2)
bullets([
    "Occupation and education can be collapsed into a smaller number of "
    "behavioural groups for campaign creative, using the embedding geometry rather "
    "than an analyst's intuition about which levels belong together.",
    "The near-zero country correlation argues against country-specific creative as "
    "a first-order lever; household composition remains the stronger axis.",
    "Because the embedding recovered the ordering of education unprompted, the "
    "ordinal rank encoding used in the benchmark feature set is validated as a "
    "reasonable simplification rather than an arbitrary modelling choice.",
])

# ==========================================================================
h("5. Conclusion and Recommendation", 1)
para("Recommendation: retain the gradient-boosting model in production.",
     bold=True)
bullets([
    "It leads on ROC-AUC (" + format(BM_AUC, ".4f") + " against "
    + format(NN_AUC, ".4f") + "), which is the metric the campaign consumes.",
    "It trains in seconds on CPU, with no GPU dependency and no epoch or "
    "learning-rate schedule to maintain.",
    "It is simpler to retrain, deploy and explain, and the accuracy difference "
    "does not justify the additional operational burden.",
])
para("Retain the embedding model as an analytical instrument, not as the scorer.")
bullets([
    "The learned embeddings are the only artefact in the project that shows which "
    "categorical levels behave alike, which is directly usable for segment design.",
    "The network reaches a materially higher recall operating point, which would "
    "be preferable if the business priority shifted from ranking efficiency to "
    "reaching as many buyers as possible.",
])
para("Where a neural approach would be expected to win:", bold=True)
bullets([
    "Substantially more data - tens or hundreds of thousands of customers rather "
    "than sixteen thousand.",
    "Higher-cardinality categorical fields such as product SKU, postcode or "
    "individual store, where one-hot encoding becomes unwieldy and embeddings "
    "compress meaningfully. StateProvinceName, dropped during feature selection "
    "for having 52 sparse levels, is exactly this kind of field and could be "
    "reintroduced as an embedding.",
    "Sequential behavioural data - browsing sessions, purchase histories, campaign "
    "response over time - where recurrent or attention-based architectures can "
    "exploit order that trees cannot.",
    "Multi-task learning - predicting purchase propensity and average monthly "
    "spend jointly from a shared representation, so the two targets regularise "
    "each other.",
])

h("Appendix A - Reproducibility", 1)
table_from_df(pd.DataFrame([
    {"Item": "Script", "Value": "aw_deep_learning.py"},
    {"Item": "Framework", "Value": D["framework"]},
    {"Item": "Device", "Value": D["device"]},
    {"Item": "Random seed", "Value": "42 (numpy and torch, including CUDA)"},
    {"Item": "Saved model", "Value": "outputs_dl/tables/dl_model.pt"},
    {"Item": "Figures", "Value": "outputs_dl/figures/ (3 plots)"},
    {"Item": "Result tables", "Value": "outputs_dl/tables/ (6 files)"},
    {"Item": "Benchmark comparison source", "Value": "outputs/tables/"},
    {"Item": "Git repository", "Value": "https://github.com/SangithaV2024BC26670/amldl-assignment-2"},
]), widths=[1.8, 4.5])

doc.save(OUT)
print("Wrote " + OUT)
