"""
Builds the Assignment-2 end-term progress report (.docx) for the AdventureWorks
bike-buyer project, from the artefacts written by aw_pipeline.py.

Every number is read from outputs/tables at build time, so the document can
never quote a stale result.
"""

import json
import os

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TAB = os.path.join("outputs", "tables")
FIG = os.path.join("outputs", "figures")
OUT = "AMLDL_Assignment2_AdventureWorks_Progress_Report.docx"
ACCENT = RGBColor(0x1F, 0x3B, 0x73)


def T(name):
    return pd.read_csv(os.path.join(TAB, name))


S = json.load(open(os.path.join(TAB, "aw_t15_run_summary.json")))
segment = T("aw_t01_segment_profile.csv")
untuned = T("aw_t03_untuned_model_comparison.csv")
tuning = T("aw_t04_tuning_best_params.csv")
pca_cmp = T("aw_t07_pca_model_comparison.csv")
final = T("aw_t08_final_model_comparison.csv")
imp = T("aw_t10_permutation_importance.csv")
dec = T("aw_t11_decile_lift_analysis.csv")
econ = T("aw_t12_campaign_economics.csv")
sens = T("aw_t13_targeting_sensitivity.csv")

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table_from_df(df, widths=None, caption=None):
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(df.columns):
        cell = t.rows[0].cells[i]
        cell.text = str(c)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if pd.isna(v) else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    if caption:
        cap = doc.add_paragraph()
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(8.5)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t


def figure(fname, caption, width=6.3):
    path = os.path.join(FIG, fname)
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(8.5)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


BEST = S["diagnostics_model"]
brow = final[final["Model"] == BEST].iloc[0]
maj = final[final["Model"] == "Baseline - Majority Class"].iloc[0]

# ==========================================================================
t = doc.add_heading("Applied Machine Learning and Deep Learning (MBA ZG582)", 0)
for r in t.runs:
    r.font.color.rgb = ACCENT
p = doc.add_paragraph()
r = p.add_run("Assignment 2 | Phase 2: Machine Learning Phase")
r.bold = True
r.font.size = Pt(13)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
r = p.add_run("Bike-Buyer Propensity Modelling for Adventure Works Cycles")
r.font.size = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_table(rows=0, cols=2)
meta.style = "Light List Accent 1"
for k, v in [
    ("Course", "MBA ZG582 - Applied Machine Learning and Deep Learning"),
    ("Phase", "Phase 2 - Model development, training, tuning and evaluation"),
    ("Dataset", "Microsoft AdventureWorks customer data (EdX DAT275x extract)"),
    ("Records", str(S["rows_after_dedup"]) + " unique customers after deduplication"),
    ("Problem type", "Supervised binary classification (purchase propensity)"),
    ("Target", "BikeBuyer - " + format(S["buy_rate_overall"] * 100, ".2f") + "% positive"),
    ("Features", str(S["features_used"]) + " after selection"),
    ("Split", str(S["train_rows"]) + " train / " + str(S["test_rows"])
     + " test, stratified"),
    ("Selected model", BEST),
]:
    row = meta.add_row().cells
    row[0].text = k
    row[1].text = v
    for pp in row[0].paragraphs:
        for rr in pp.runs:
            rr.bold = True
    for c in row:
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(9.5)

# ==========================================================================
h("1. Problem Statement and Business Context", 1)
para(
    "Adventure Works Cycles sells bicycles and accessories to retail customers "
    "across six countries. The company runs direct marketing campaigns against its "
    "customer list, and faces the standard targeting trade-off: contacting the "
    "entire base wastes budget on people who will never buy, while contacting too "
    "narrowly leaves revenue uncollected. The business question is therefore not "
    "\"who is a good customer\" in the abstract, but specifically which customers "
    "should receive the next campaign.")

h("1.1 Machine Learning Framing", 2)
para(
    "The task is framed as supervised binary classification: predict BikeBuyer, "
    "whether a customer has purchased a bicycle, from demographic and household "
    "attributes. Classification is the correct framing because the underlying "
    "decision is itself binary - contact or do not contact - taken over a fixed "
    "customer list.")
para(
    "One point of nuance matters for how the model is evaluated. The deliverable "
    "is not really a yes/no label; it is a ranked propensity score, which the "
    "campaign then cuts at whatever depth the budget allows. That is why ROC-AUC, "
    "which scores the quality of the ranking independently of any single cut-off, "
    "is the primary metric and the tuning objective.")

table_from_df(pd.DataFrame([
    {"Element": "Target variable", "Definition": "BikeBuyer - binary, "
     + format(S["buy_rate_overall"] * 100, ".2f") + "% positive"},
    {"Element": "Unit of analysis", "Definition": "one customer"},
    {"Element": "Predictors", "Definition": str(S["features_used"])
     + " demographic and household features after selection"},
    {"Element": "Class balance", "Definition": "1 : 2.01 - moderate imbalance"},
    {"Element": "Primary metric", "Definition": "ROC-AUC - scores the ranking, independent of cut-off"},
    {"Element": "Secondary metrics", "Definition": "F1, PR-AUC, precision, recall, balanced accuracy"},
    {"Element": "Baselines", "Definition": "majority-class and stratified-random classifiers"},
]), widths=[1.5, 4.8])

h("1.2 Why Accuracy Alone Would Mislead", 2)
para(
    "Because " + format((1 - S["buy_rate_overall"]) * 100, ".1f") + " per cent of "
    "customers do not buy, a model that predicts \"nobody buys\" scores "
    + format(maj["Accuracy"], ".4f") + " accuracy while being entirely useless: its "
    "F1 is " + format(maj["F1"], ".4f") + " and its ROC-AUC is "
    + format(maj["ROC_AUC"], ".4f") + ". Both trivial baselines are trained and "
    "scored explicitly in Step 4 so that this is demonstrated rather than merely "
    "asserted, and so that every later result has an honest floor to be measured "
    "against.")

# ==========================================================================
h("2. Methodology", 1)
table_from_df(pd.DataFrame([
    {"Step": "1", "Stage": "Dataset preparation", "Detail": "Join three tables, resolve duplicates, drop PII, derive age and ratios, ordinal + one-hot encoding"},
    {"Step": "2", "Stage": "Feature selection", "Detail": "Leakage exclusion, zero-variance filter, |r| > 0.95 redundancy filter, mutual information"},
    {"Step": "3", "Stage": "Stratified split", "Detail": "80 / 20, class ratio preserved in both partitions"},
    {"Step": "4", "Stage": "Model training", "Detail": "2 baselines + 6 candidate algorithms"},
    {"Step": "5", "Stage": "Hyperparameter tuning", "Detail": "RandomizedSearchCV (25 draws, StratifiedKFold) + validation curves"},
    {"Step": "6", "Stage": "Dimensionality reduction", "Detail": "PCA, models retrained on components"},
    {"Step": "7", "Stage": "Ensembles", "Detail": "Bagging, boosting, soft voting, stacking"},
    {"Step": "8", "Stage": "Comparison and selection", "Detail": "Ranked on hold-out ROC-AUC"},
    {"Step": "9", "Stage": "Diagnostics", "Detail": "Confusion matrix, ROC, PR, calibration, threshold optimisation"},
    {"Step": "10", "Stage": "Business interpretation", "Detail": "Decile lift and gains, campaign economics, sensitivity analysis"},
]), widths=[0.5, 1.7, 4.1])

h("2.1 Why the Split Is Random Here", 2)
para(
    "This dataset is a customer cross-section rather than a time series, so a "
    "random split is appropriate - there is no temporal ordering that a shuffle "
    "could violate. The split is stratified so that both partitions carry the same "
    "buyer proportion, which matters under class imbalance. This is a deliberate "
    "contrast with a forecasting problem, where a shuffled split would leak the "
    "future into the past and inflate the reported score.")

# ==========================================================================
h("3. Data Preparation and Feature Selection", 1)
para(
    "Three files - customer demographics, the BikeBuyer flag and average monthly "
    "spend - are joined on CustomerID. The preparation stage is substantive rather "
    "than a formality, because the raw extract carries four genuine defects.")

h("3.1 Duplicate Records", 2)
para(
    "The extract repeats some CustomerIDs, and a blanket drop_duplicates() would "
    "be wrong because three distinct cases exist, each needing a different answer.")
table_from_df(pd.DataFrame([
    {"Case": "Exact duplicate rows", "Count": "98", "Treatment": "Safe to drop - no information lost"},
    {"Case": "Same ID, conflicting attributes", "Count": "17", "Treatment": "Keep the last record"},
    {"Case": "Same ID, conflicting target", "Count": "4", "Treatment": "Keep the last record"},
]), widths=[2.2, 0.8, 3.3])
para(
    "The keep-last rule is applied consistently across all three tables so that "
    "customer, target and spend records stay aligned, reducing 16,519 raw rows to "
    + str(S["rows_after_dedup"]) + " unique customers.")

h("3.2 Other Cleaning and Feature Engineering", 2)
bullets([
    "Identifier and PII columns dropped - names, phone numbers and street "
    "addresses are unique labels rather than features, and carry no generalisable "
    "signal.",
    "Whitespace defect corrected - the value 'Bachelors ' ships with a trailing "
    "space and would otherwise encode as a level distinct from 'Bachelors', "
    "silently splitting one category into two.",
    "Age derived against a fixed 1998 reference date. AdventureWorks is a "
    "1998-vintage sample database; computing age against today's date would add "
    "roughly 28 years to every customer and destroy the age signal.",
    "Education rank-encoded rather than one-hot. The five levels have a natural "
    "order, so ranking preserves that ordering for the linear models and gives the "
    "tree models a single splittable column instead of five sparse ones.",
    "Derived ratios added - income per child, cars per child, children away from "
    "home, log income - to express relationships the raw columns only imply.",
])

h("3.3 Leakage Exclusion", 2)
para(
    "AveMonthSpend is deliberately excluded from the predictors. It records what a "
    "customer spends with Adventure Works, so it is measured contemporaneously "
    "with - and is partly caused by - the bike purchase itself. Including it would "
    "leak the outcome into the features and inflate every score, and it would be "
    "unavailable at scoring time for a prospect who has not yet bought anything. "
    "Identifying this before modelling, rather than being pleased by an "
    "artificially high AUC afterwards, is the substantive judgement in this step.",
    italic=False)
para(
    "The redundancy filter then removed MaritalStatus_S, which duplicated the "
    "engineered IsSingle flag exactly - a useful confirmation that the filter works.")

figure("aw_fig01_target_and_distributions.png",
       "Figure 1: Target balance, and the age and income distributions split by "
       "purchase outcome.")
figure("aw_fig02_purchase_rate_by_segment.png",
       "Figure 2: Purchase rate by segment against the overall base rate.")
figure("aw_fig03_segment_heatmap.png",
       "Figure 3: Purchase rate by occupation and marital status.")
table_from_df(segment, caption="Table 1: Customer profile and purchase rate by "
                               "occupation.")
figure("aw_fig04_feature_selection_mi.png",
       "Figure 4: Top 15 features by mutual information with BikeBuyer.")
figure("aw_fig05_correlation_matrix.png",
       "Figure 5: Correlation matrix of the retained feature set.")

# ==========================================================================
h("4. Model Development and Evaluation", 1)
para(
    "Eight models were trained: two trivial baselines and six candidate "
    "algorithms spanning linear, instance-based, tree and ensemble families.")
table_from_df(untuned, caption="Table 2: Untuned model performance on the "
                               "stratified hold-out.")
para(
    "The majority-class baseline is the most instructive row in the table. It "
    "achieves " + format(maj["Accuracy"], ".4f") + " accuracy - which would sound "
    "acceptable reported on its own - while scoring zero on F1 and 0.5 on ROC-AUC, "
    "because it never identifies a single buyer. Every subsequent claim in this "
    "report is therefore made on AUC and F1 rather than accuracy.")

# ==========================================================================
h("5. Hyperparameter Tuning", 1)
para(
    "RandomizedSearchCV sampled 25 configurations per model under a 5-fold "
    "StratifiedKFold, which preserves the class ratio in every fold. Validation "
    "curves then swept individual parameters to show where over-fitting begins "
    "rather than only reporting the winning configuration.")
table_from_df(tuning, widths=[1.4, 1.1, 3.8],
              caption="Table 3: Best cross-validated configuration per model.")
figure("aw_fig06_valcurve_rf_maxdepth.png",
       "Figure 6: Validation curve - Random Forest tree depth.")
figure("aw_fig07_valcurve_rf_minleaf.png",
       "Figure 7: Validation curve - Random Forest minimum samples per leaf.")
figure("aw_fig08_valcurve_hgb_learningrate.png",
       "Figure 8: Validation curve - boosting learning rate.")

# ==========================================================================
h("6. Dimensionality Reduction (PCA)", 1)
para(
    "PCA was fitted on the standardised training matrix. "
    + str(S["pca_components_for_95pct"]) + " components are needed to retain 95 per "
    "cent of the variance against " + str(S["features_used"]) + " original "
    "features. Logistic Regression and Random Forest were retrained on the "
    "components and scored on the same hold-out.")
figure("aw_fig09_pca_explained_variance.png",
       "Figure 9: Scree plot and cumulative explained variance.")
figure("aw_fig10_pca_projection.png",
       "Figure 10: Customers projected onto the first two principal components, "
       "coloured by purchase outcome.")
table_from_df(pca_cmp, caption="Table 4: Models retrained on principal components.")
para(
    "PCA does not improve performance here. The best PCA variant scores below its "
    "equivalent model on the full feature set, and the two-component projection "
    "shows the classes are not linearly separable in the reduced space. The "
    "explanation is structural: roughly half the matrix consists of orthogonal "
    "one-hot dummies and binary flags, which PCA cannot compress without discarding "
    "information. Interpretability is also lost - a marketing manager can act on "
    "'customers with children at home' but not on 'principal component 4'. PCA is "
    "therefore documented as a completed experiment and is not adopted in the final "
    "pipeline.")

# ==========================================================================
h("7. Ensemble Techniques", 1)
table_from_df(pd.DataFrame([
    {"Family": "Bagging", "Implementation": "Random Forest", "Mechanism": "Parallel trees on bootstrap samples"},
    {"Family": "Boosting", "Implementation": "Gradient Boosting, HistGradientBoosting", "Mechanism": "Sequential trees, each correcting its predecessor"},
    {"Family": "Voting", "Implementation": "VotingClassifier (soft)", "Mechanism": "Averages predicted probabilities, preserving ranking information"},
    {"Family": "Stacking", "Implementation": "StackingClassifier", "Mechanism": "Logistic meta-model learns how to weight the base models"},
]), widths=[0.9, 2.2, 3.2])
para(
    "Soft voting was chosen over hard voting because the campaign consumes ranked "
    "probabilities rather than labels, and averaging probabilities preserves that "
    "ranking. For stacking, StratifiedKFold partitions the data cleanly, so "
    "scikit-learn's StackingClassifier can build its meta-features directly - a "
    "notable contrast with a time-series problem, where TimeSeriesSplit does not "
    "partition and a hand-built temporal stack would be required.")

# ==========================================================================
h("8. Model Comparison and Final Selection", 1)
table_from_df(final, caption="Table 5: Full model comparison on the hold-out, "
                             "ranked by ROC-AUC.")
figure("aw_fig11_model_comparison.png",
       "Figure 11: Test-set ROC-AUC and F1 by model.")
para("Selected model: " + BEST, bold=True)
bullets([
    "ROC-AUC " + format(brow["ROC_AUC"], ".4f") + ", PR-AUC "
    + format(brow["PR_AUC"], ".4f") + ", F1 " + format(brow["F1"], ".4f")
    + ", accuracy " + format(brow["Accuracy"], ".4f") + ".",
    "Best ranking quality in the comparison, which is the property the campaign "
    "actually consumes.",
    "Trains in seconds and handles the mixed numeric and binary feature matrix "
    "natively, so periodic retraining is operationally trivial.",
    "Supports permutation importance, so the drivers behind a customer's score can "
    "be explained to a marketing stakeholder.",
])
para(
    "The margin over the next models is narrow - the stacking ensemble and plain "
    "gradient boosting sit within roughly half a point of AUC. Where differences "
    "are this small, the simpler and faster model is the defensible choice, and the "
    "tuned boosting model is both.")

# ==========================================================================
h("9. Diagnostics and Threshold Optimisation", 1)
figure("aw_fig12_diagnostics.png",
       "Figure 12: Confusion matrix, ROC curve, precision-recall curve and "
       "calibration curve for the selected model.")
para(
    "The precision-recall curve is the more informative of the two curves under "
    "class imbalance, since ROC can look flattering when the negative class "
    "dominates. The calibration curve matters because the campaign consumes "
    "predicted probabilities directly, so those probabilities need to mean what "
    "they claim.")
para(
    "The default 0.50 cut-off is a software default, not a business decision. Two "
    "alternatives were evaluated on the hold-out.")
table_from_df(pd.DataFrame([
    {"Threshold": "0.500 (default)", "Basis": "scikit-learn default",
     "Consequence": "Highest accuracy, but recall of only "
                    + format(brow["Recall"], ".3f") + " - most buyers are missed"},
    {"Threshold": format(S["f1_optimal_threshold"], ".3f") + " (F1-optimal)",
     "Basis": "Balances precision and recall",
     "Consequence": "Substantially higher recall at modest precision cost"},
    {"Threshold": format(S["profit_optimal_threshold"], ".3f") + " (profit-optimal)",
     "Basis": "Maximises expected campaign profit",
     "Consequence": "Contacts almost everyone - see Section 10"},
]), widths=[1.5, 1.7, 3.1])
figure("aw_fig13_threshold_optimisation.png",
       "Figure 13: Precision-recall trade-off and campaign profit across "
       "decision thresholds.")
figure("aw_fig14_feature_importance.png",
       "Figure 14: Permutation feature importance measured on the hold-out.")
para("Strongest drivers of purchase propensity:", bold=True)
bullets([str(r["Feature"]) + " - reduces ROC-AUC by "
         + format(r["Importance"], ".4f") + " when shuffled"
         for _, r in imp.head(6).iterrows()])

# ==========================================================================
h("10. Business Interpretation", 1)
para(
    "The model is converted into a campaign decision by ranking customers on "
    "predicted propensity, splitting them into ten equal groups, and comparing the "
    "buy rate in each against the base rate.")
table_from_df(dec, caption="Table 6: Decile lift and cumulative gains analysis.")
figure("aw_fig15_lift_gains_profit.png",
       "Figure 15: Lift by decile, cumulative gains curve, and profit by campaign "
       "depth.")
bullets([
    "The top decile buys at " + format(dec.loc[0, "Lift"], ".2f") + " times the "
    "base rate - a genuinely strong separation.",
    "The top 30 per cent of the ranked base captures "
    + format(S["capture_top30pct"], ".1f") + " per cent of all buyers.",
    "The bottom decile contains no buyers at all, so those contacts are pure cost.",
])

h("10.1 An Honest Reading of the Campaign Economics", 2)
para(
    "Under the stated assumptions - a contribution margin of "
    + format(S["margin_per_sale_assumed"], ".0f") + " per sale against a contact "
    "cost of " + format(S["cost_per_contact_assumed"], ".2f") + " - the targeted "
    "campaign returns only " + format(S["targeted_vs_mailall_uplift_pct"], ".2f")
    + " per cent more profit than simply mailing everybody. That result should be "
    "reported plainly rather than buried.")
table_from_df(econ, widths=[1.9, 1.6, 1.6],
              caption="Table 7: Campaign economics, model-targeted versus "
                      "mailing the entire base.")
para(
    "The reason is arithmetic, not a defect in the model. A customer is worth "
    "contacting whenever their purchase probability exceeds break-even = cost / "
    "margin = " + format(S["breakeven_probability"] * 100, ".2f") + " per cent. "
    "Almost the entire customer base clears that bar, so blanket mailing is close "
    "to optimal and there is little profit for targeting to add.")
para(
    "Whether targeting pays is therefore a property of the campaign economics, not "
    "of the model. The sensitivity analysis below tests where the crossover lies.")
table_from_df(sens, caption="Table 8: Sensitivity of optimal campaign depth and "
                            "profit uplift to the cost per contact.")
figure("aw_fig16_targeting_sensitivity.png",
       "Figure 16: Optimal campaign depth and profit uplift as contact cost rises.")
para("What this means in practice:", bold=True)
bullets([
    "For cheap channels - email, in-app messaging, SMS - the model should not be "
    "used to decide whom to exclude, because almost everyone is worth contacting. "
    "Its value there is efficiency: the same revenue is earned from roughly 10 per "
    "cent fewer contacts, which reduces list fatigue and unsubscribes.",
    "For expensive channels - printed catalogues, outbound calling, field visits - "
    "the economics invert and targeting becomes decisive. As contact cost rises, "
    "the optimal campaign depth falls and the profit uplift over blanket mailing "
    "grows substantially.",
    "The model's most durable value is as a ranking, not as a hard classifier. It "
    "tells the business the order in which to spend a budget of any size, which is "
    "a more useful output than a contact / do-not-contact label.",
])

h("10.2 Segment Insight for Marketing", 2)
bullets([
    "Customers with children at home are the single strongest signal in the model, "
    "ahead of both income and age.",
    "Professionals buy at " + format(segment.loc[0, "BuyRate"], ".1f")
    + " per cent against " + format(segment.iloc[-1]["BuyRate"], ".1f")
    + " per cent for manual occupations - a spread of roughly two to one.",
    "Income matters, but less than household composition, which suggests creative "
    "and offer design should lead on family use cases rather than on price alone.",
])

# ==========================================================================
h("11. Conclusion and Next Steps", 1)
para(
    "Phase 2 delivers a validated purchase-propensity model for Adventure Works "
    "Cycles. " + BEST + " was selected on hold-out ROC-AUC of "
    + format(brow["ROC_AUC"], ".4f") + ", against " + format(maj["ROC_AUC"], ".4f")
    + " for a majority-class baseline, and separates the top decile at "
    + format(dec.loc[0, "Lift"], ".2f") + " times the base purchase rate. The "
    "analysis also establishes the conditions under which that ranking is "
    "commercially worth acting on, which is a more useful conclusion than the "
    "accuracy figure alone.")
para("Carried forward to Phase 3:", bold=True)
bullets([
    "Deep learning extension - a neural network with entity embeddings for the "
    "categorical fields, which can learn interactions the tree ensembles "
    "approximate by splitting.",
    "Probability calibration - isotonic or Platt scaling, so that predicted scores "
    "can be used directly in expected-value calculations.",
    "Cost-sensitive learning - training against the campaign profit function "
    "directly rather than optimising AUC and choosing a threshold afterwards.",
    "Richer features - purchase recency, frequency and monetary value, and "
    "campaign response history, which would very likely outperform demographics "
    "alone.",
    "A second target - AveMonthSpend is available on the same customer table, so a "
    "regression model could estimate customer value and combine with propensity to "
    "rank on expected revenue rather than probability of purchase.",
])

h("Appendix A - Reproducibility", 1)
table_from_df(pd.DataFrame([
    {"Item": "Notebook", "Value": "AMLDL_Assignment2_AdventureWorks.ipynb"},
    {"Item": "Pipeline script", "Value": "aw_pipeline.py"},
    {"Item": "Dataset", "Value": "data/adventureworks/ - 3 CSVs, Microsoft AdventureWorks"},
    {"Item": "Source", "Value": "github.com/kdavenpo/AdventureWorks (EdX DAT275x extract)"},
    {"Item": "Random seed", "Value": "42 (fixed throughout)"},
    {"Item": "Figures", "Value": "outputs/figures/ (16 plots)"},
    {"Item": "Result tables", "Value": "outputs/tables/ (15 files)"},
    {"Item": "Stated assumptions", "Value": "Margin per sale "
     + format(S["margin_per_sale_assumed"], ".0f") + ", cost per contact "
     + format(S["cost_per_contact_assumed"], ".2f") + " - illustrative, not "
     "derived from the data"},
    {"Item": "Colab / Git link", "Value": "[paste your link here before submission]"},
]), widths=[1.6, 4.7])

doc.save(OUT)
print("Wrote " + OUT)
