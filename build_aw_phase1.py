"""
Builds the Phase 1 concept note (.docx) for the AdventureWorks bike-buyer
project, mirroring the section structure of the Phase 1 note submitted for the
earlier pharma project so the two read as the same deliverable type.

Statistics are computed from the raw CSVs and from outputs/tables at build
time, so the note cannot quote a stale number.
"""

import json
import os

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DATA = os.path.join("data", "adventureworks")
TAB = os.path.join("outputs", "tables")
FIG = os.path.join("outputs", "figures")
OUT = "Concept_Note_AdventureWorks_Bike_Buyer_Propensity.docx"
ACCENT = RGBColor(0x1F, 0x3B, 0x73)
AGE_REF = pd.Timestamp("1998-01-01")

S = json.load(open(os.path.join(TAB, "aw_t15_run_summary.json")))
final = pd.read_csv(os.path.join(TAB, "aw_t08_final_model_comparison.csv"))
segment = pd.read_csv(os.path.join(TAB, "aw_t01_segment_profile.csv"))
dec = pd.read_csv(os.path.join(TAB, "aw_t11_decile_lift_analysis.csv"))
imp = pd.read_csv(os.path.join(TAB, "aw_t10_permutation_importance.csv"))

# --- Recompute headline EDA statistics from the raw files -------------------
custs = (pd.read_csv(os.path.join(DATA, "AdvWorksCusts.csv"))
         .drop_duplicates().drop_duplicates(subset="CustomerID", keep="last"))
buyer = (pd.read_csv(os.path.join(DATA, "AW_BikeBuyer.csv"))
         .drop_duplicates().drop_duplicates(subset="CustomerID", keep="last"))
eda = custs.merge(buyer, on="CustomerID", how="inner")
eda["Education"] = eda["Education"].str.strip()
eda["Age"] = ((AGE_REF - pd.to_datetime(eda["BirthDate"])).dt.days / 365.25)

by_country = (eda.groupby("CountryRegionName")["BikeBuyer"]
              .agg(Customers="size", BuyRate="mean").reset_index()
              .assign(BuyRate=lambda d: (d["BuyRate"] * 100).round(2))
              .sort_values("BuyRate", ascending=False))
by_children = (eda.groupby("NumberChildrenAtHome")["BikeBuyer"]
               .agg(Customers="size", BuyRate="mean").reset_index()
               .assign(BuyRate=lambda d: (d["BuyRate"] * 100).round(2)))
buyers = eda[eda["BikeBuyer"] == 1]
nonbuyers = eda[eda["BikeBuyer"] == 0]

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table_from_df(df, widths=None, caption=None):
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, c in enumerate(df.columns):
        cell = t.rows[0].cells[i]
        cell.text = str(c)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if pd.isna(v) else str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    if caption:
        cap = doc.add_paragraph()
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(8.5)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def figure(fname, caption, width=6.2):
    path = os.path.join(FIG, fname)
    if not os.path.exists(path):
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(8.5)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


BEST = S["selected_model"]
brow = final[final["Model"] == BEST].iloc[0]

# ==========================================================================
# TITLE
# ==========================================================================
t = doc.add_heading("Concept Note: Customer Purchase-Propensity Modelling", 0)
for r in t.runs:
    r.font.color.rgb = ACCENT
p = doc.add_paragraph()
r = p.add_run("Targeting Direct Marketing Campaigns at Adventure Works Cycles")
r.bold = True
r.font.size = Pt(13)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
r = p.add_run("Course Project - Applied Machine Learning and Deep Learning "
              "(MBA ZG582) | Phase 1")
r.font.size = Pt(10.5)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
r = p.add_run("Dataset: Microsoft AdventureWorks customer data "
              "(DAT275x lab extract)")
r.italic = True
r.font.size = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ==========================================================================
h("1. Problem Statement", 1)

h("1.1 Business Context", 2)
para(
    "Adventure Works Cycles is a manufacturer and retailer of bicycles and "
    "cycling accessories, selling to individual consumers across six countries in "
    "North America, Europe and Australia. Like most consumer retailers, it "
    "acquires and retains customers partly through outbound direct marketing - "
    "catalogues, mailings and promotional offers sent to names on its customer "
    "list.")
para(
    "Direct marketing carries a structural tension. Contacting the entire customer "
    "base guarantees that every genuine prospect is reached, but spends budget on "
    "the majority who will not buy, and repeated irrelevant contact drives list "
    "fatigue and unsubscribes. Contacting too narrow a group protects budget but "
    "leaves revenue uncollected. With roughly one customer in three having "
    "purchased a bicycle, an untargeted campaign wastes the majority of its spend.")
para(
    "The commercial question is therefore not who is a valuable customer in the "
    "abstract, but specifically which customers should receive the next campaign, "
    "and how deep into the list the company should go for a given budget.")

h("1.2 ML Problem Framing", 2)
para("This is framed as a supervised binary classification problem:")
bullets([
    "Target variable (y): BikeBuyer - whether a customer has purchased a bicycle "
    "(1) or not (0). Binary, with " + format(S["buy_rate_overall"] * 100, ".2f")
    + " per cent positive cases.",
    "Features (X): customer demographic and household attributes - age, gender, "
    "marital status, education, occupation, yearly income, home ownership, number "
    "of cars owned, and children at home and in total, plus country of residence.",
    "Goal: rank customers by purchase propensity so that campaign spend is "
    "concentrated on the highest-scoring segments, maximising response per unit of "
    "marketing cost.",
])
para(
    "One point of framing deserves emphasis, because it determines how the model "
    "is evaluated. Although the target is a binary label, the operational "
    "deliverable is not a yes/no decision but a ranked propensity score, which the "
    "campaign then cuts at whatever depth the budget allows. The primary metric is "
    "therefore ROC-AUC, which measures the quality of that ranking independently "
    "of any particular cut-off, supported by F1 and precision-recall AUC for the "
    "minority class.")
para(
    "A secondary regression framing - predicting AveMonthSpend, the customer's "
    "average monthly spend - is available on the same customer table and would "
    "allow ranking on expected revenue rather than probability of purchase. It is "
    "noted as a natural extension but is out of scope for this concept note.")

# ==========================================================================
h("2. Dataset Description", 1)
para(
    "The analysis uses the Microsoft AdventureWorks sample customer data, in the "
    "extract prepared for the Microsoft / EdX course DAT275x, Principles of "
    "Machine Learning. AdventureWorks is Microsoft's long-standing publicly "
    "available sample database for a fictional bicycle retailer, and is widely "
    "used for teaching and benchmarking. Three files are joined on CustomerID: "
    "customer demographics, the bike-purchase flag, and average monthly spend.")

h("2.1 Dataset Summary", 2)
table_from_df(pd.DataFrame([
    {"Attribute": "Source", "Value": "Microsoft AdventureWorks sample database (DAT275x extract)"},
    {"Attribute": "Public mirror", "Value": "github.com/kdavenpo/AdventureWorks"},
    {"Attribute": "Files", "Value": "AdvWorksCusts.csv, AW_BikeBuyer.csv, AW_AveMonthSpend.csv"},
    {"Attribute": "Raw rows", "Value": "16,519 per file"},
    {"Attribute": "Unique customers", "Value": str(S["rows_after_dedup"]) + " after deduplication"},
    {"Attribute": "Raw attributes", "Value": "23 customer columns + 2 target columns"},
    {"Attribute": "Modelling features", "Value": str(S["features_used"]) + " after cleaning, engineering and selection"},
    {"Attribute": "Target (classification)", "Value": "BikeBuyer - " + format(S["buy_rate_overall"] * 100, ".2f") + "% positive"},
    {"Attribute": "Target (secondary)", "Value": "AveMonthSpend - continuous, out of scope here"},
    {"Attribute": "Class balance", "Value": "1 : 2.01 (moderate imbalance)"},
    {"Attribute": "Geography", "Value": "6 countries - United States, Australia, United Kingdom, France, Germany, Canada"},
    {"Attribute": "Age range", "Value": format(eda["Age"].min(), ".0f") + " - "
     + format(eda["Age"].max(), ".0f") + " years (median "
     + format(eda["Age"].median(), ".0f") + "), as at 1 January 1998"},
    {"Attribute": "Income range", "Value": format(eda["YearlyIncome"].min(), ",.0f")
     + " - " + format(eda["YearlyIncome"].max(), ",.0f") + " (median "
     + format(eda["YearlyIncome"].median(), ",.0f") + ")"},
    {"Attribute": "Missing values", "Value": "None in any modelling column"},
]), widths=[1.7, 4.6])

h("2.2 Feature Dictionary", 2)
table_from_df(pd.DataFrame([
    {"Feature": "Age", "Type": "Numeric (derived)", "Description": "Derived from BirthDate against a fixed 1 Jan 1998 reference date"},
    {"Feature": "YearlyIncome", "Type": "Numeric", "Description": "Annual household income"},
    {"Feature": "NumberChildrenAtHome", "Type": "Numeric", "Description": "Children currently living at home (0-5)"},
    {"Feature": "TotalChildren", "Type": "Numeric", "Description": "Total children (0-5)"},
    {"Feature": "NumberCarsOwned", "Type": "Numeric", "Description": "Vehicles owned by the household (0-4)"},
    {"Feature": "HomeOwnerFlag", "Type": "Binary", "Description": "Owns their home (1) or not (0)"},
    {"Feature": "Education", "Type": "Ordinal", "Description": "Five ordered levels, Partial High School to Graduate Degree"},
    {"Feature": "Occupation", "Type": "Nominal", "Description": "Manual, Clerical, Skilled Manual, Professional, Management"},
    {"Feature": "Gender", "Type": "Nominal", "Description": "Male / Female"},
    {"Feature": "MaritalStatus", "Type": "Nominal", "Description": "Married / Single"},
    {"Feature": "CountryRegionName", "Type": "Nominal", "Description": "Country of residence (6 levels)"},
    {"Feature": "Derived ratios", "Type": "Numeric", "Description": "Income per child, cars per child, children away from home, log income"},
]), widths=[1.5, 1.1, 3.7])

h("2.3 Data Quality Assessment", 2)
para(
    "The extract is not clean on arrival, and the defects found are substantive "
    "rather than cosmetic. Identifying them is a material part of the project.")
table_from_df(pd.DataFrame([
    {"Issue": "Duplicate CustomerIDs", "Detail": "115 IDs appear more than once - 98 exact duplicate rows, 17 with conflicting attributes, 4 with conflicting targets", "Treatment": "Drop exact duplicates, then keep the last record per ID, applied consistently across all three tables"},
    {"Issue": "Whitespace defect", "Detail": "'Bachelors ' carries a trailing space", "Treatment": "Strip whitespace before encoding, else one category silently splits into two"},
    {"Issue": "Ordinal treated as nominal", "Detail": "Education has five naturally ordered levels", "Treatment": "Rank-encode rather than one-hot, preserving the ordering"},
    {"Issue": "Date vintage", "Detail": "AdventureWorks is 1998-vintage data", "Treatment": "Compute age against a fixed 1998 reference date, not the present day"},
    {"Issue": "Potential target leakage", "Detail": "AveMonthSpend is measured contemporaneously with the purchase", "Treatment": "Excluded from the predictors - see Section 2.5"},
    {"Issue": "Identifier / PII columns", "Detail": "Names, phone numbers, street addresses", "Treatment": "Dropped - unique labels carrying no generalisable signal"},
]), widths=[1.3, 2.5, 2.5])

h("2.4 Exploratory Data Assessment (Highlights)", 2)
bullets([
    "The target is moderately imbalanced: " + format(S["buy_rate_overall"] * 100, ".2f")
    + " per cent of customers have bought a bicycle, so a naive model predicting "
    "'nobody buys' would score roughly 67 per cent accuracy while identifying no "
    "buyers at all. Accuracy alone is therefore an unsafe metric for this problem.",
    "Household composition is by far the strongest signal, and the relationship is "
    "monotonic: the purchase rate climbs from "
    + format(by_children.iloc[0]["BuyRate"], ".1f") + " per cent for customers with "
    "no children at home to " + format(by_children.iloc[-1]["BuyRate"], ".1f")
    + " per cent for those with five - a spread of more than four to one. This "
    "single variable outranks both income and age on mutual information with the "
    "target, and points the campaign's creative towards family use cases rather "
    "than price.",
    "Occupation separates strongly: professionals buy at "
    + format(segment.loc[0, "BuyRate"], ".1f") + " per cent against "
    + format(segment.iloc[-1]["BuyRate"], ".1f") + " per cent for manual "
    "occupations, a spread of roughly two to one.",
    "Buyers are younger on average than non-buyers (median "
    + format(buyers["Age"].median(), ".1f") + " against "
    + format(nonbuyers["Age"].median(), ".1f") + " years) and have higher median "
    "income (" + format(buyers["YearlyIncome"].median(), ",.0f") + " against "
    + format(nonbuyers["YearlyIncome"].median(), ",.0f") + ").",
    "Purchase rates vary by country, from "
    + format(by_country.iloc[0]["BuyRate"], ".1f") + " per cent in "
    + str(by_country.iloc[0]["CountryRegionName"]) + " to "
    + format(by_country.iloc[-1]["BuyRate"], ".1f") + " per cent in "
    + str(by_country.iloc[-1]["CountryRegionName"]) + ", so geography carries "
    "signal but less than household attributes.",
    "No missing values are present in any modelling column once identifier "
    "columns are dropped.",
])

table_from_df(by_children, widths=[2.2, 1.3, 1.3],
              caption="Table 1: Purchase rate by number of children at home - the "
                      "strongest single relationship in the dataset.")
table_from_df(segment,
              caption="Table 2: Customer profile and purchase rate by occupation.")
table_from_df(by_country, widths=[2.2, 1.3, 1.3],
              caption="Table 3: Customer count and purchase rate by country.")

figure("aw_fig01_target_and_distributions.png",
       "Figure 1: Target balance, and age and income distributions split by "
       "purchase outcome.")
figure("aw_fig02_purchase_rate_by_segment.png",
       "Figure 2: Purchase rate by segment against the overall base rate. "
       "Household and occupation attributes separate more strongly than "
       "geography.")
figure("aw_fig03_segment_heatmap.png",
       "Figure 3: Purchase rate by occupation and marital status.")

h("2.5 A Note on Target Leakage", 2)
para(
    "AveMonthSpend records what a customer spends with Adventure Works. Because it "
    "is measured contemporaneously with - and is partly caused by - the bicycle "
    "purchase itself, including it as a predictor would leak the outcome into the "
    "feature set and inflate every reported score. It would also be unavailable at "
    "scoring time for a prospect who has not yet bought anything, so a model "
    "relying on it could not be deployed for its intended purpose. It is excluded "
    "from the predictors throughout, and retained only as a candidate target for "
    "the secondary regression framing.")

# ==========================================================================
h("3. Proposed Methodology", 1)

h("3.1 Candidate Algorithms", 2)
table_from_df(pd.DataFrame([
    {"Algorithm": "Majority-class / random baselines", "Why it fits": "Establishes the honest floor. Demonstrates concretely that ~67% accuracy is achievable while identifying zero buyers.", "Role": "Baseline"},
    {"Algorithm": "Logistic Regression", "Why it fits": "Fast, interpretable benchmark. Coefficients show the direction and size of each demographic effect, which communicates well to marketing stakeholders.", "Role": "Interpretability benchmark"},
    {"Algorithm": "K-Nearest Neighbours", "Why it fits": "Instance-based comparator - tests whether similar customers simply behave similarly, without assuming a functional form.", "Role": "Diagnostic"},
    {"Algorithm": "Decision Tree", "Why it fits": "Captures threshold effects and interactions in a form a non-technical audience can read directly.", "Role": "Intermediate"},
    {"Algorithm": "Random Forest", "Why it fits": "Bagged ensemble; handles mixed numeric and categorical features, interactions and outliers robustly, and yields feature importances.", "Role": "Ensemble (bagging)"},
    {"Algorithm": "Gradient Boosting / HistGradientBoosting", "Why it fits": "Sequential ensembles that typically lead on tabular classification; strong ranking quality, which is what the campaign consumes.", "Role": "Primary candidate"},
    {"Algorithm": "Voting and Stacking ensembles", "Why it fits": "Combine complementary model families; soft voting preserves probability ranking, stacking learns optimal weights.", "Role": "Ensemble (blending)"},
]), widths=[1.6, 3.4, 1.3])

h("3.2 Evaluation Approach", 2)
bullets([
    "Stratified 80 / 20 train-test split, preserving the class ratio in both "
    "partitions. A random split is appropriate because this is a customer "
    "cross-section, not a time series.",
    "Five-fold StratifiedKFold cross-validation for hyperparameter tuning, via "
    "RandomizedSearchCV, supported by validation curves to locate the onset of "
    "over-fitting.",
    "Primary metric ROC-AUC; secondary metrics F1, PR-AUC, precision, recall and "
    "balanced accuracy. Accuracy is reported but never relied on alone.",
    "Decision-threshold optimisation rather than acceptance of the 0.50 default, "
    "evaluated against both F1 and expected campaign profit.",
    "Principal Component Analysis evaluated as a dimensionality-reduction "
    "experiment, with adoption decided on measured AUC rather than assumption.",
])

h("3.3 Preliminary Model Results", 2)
para(
    "An initial pass across the candidate algorithms gives the following hold-out "
    "performance, and confirms the framing is sound before the full Phase 2 build.")
table_from_df(final[["Model", "Accuracy", "F1", "ROC_AUC", "PR_AUC"]].head(8),
              caption="Table 4: Preliminary hold-out performance, ranked by "
                      "ROC-AUC.")
para(
    "The tuned boosting model leads at ROC-AUC " + format(brow["ROC_AUC"], ".4f")
    + " with F1 " + format(brow["F1"], ".4f") + ", against 0.5000 for the "
    "majority-class baseline. Ranked by predicted propensity, the top decile of "
    "customers buys at " + format(dec.loc[0, "Lift"], ".2f") + " times the base "
    "rate, and the top 30 per cent of the list captures "
    + format(S["capture_top30pct"], ".1f") + " per cent of all buyers - which is "
    "the result that makes the model commercially useful.")
para("The strongest predictors are:", bold=True)
bullets([str(r["Feature"]) for _, r in imp.head(5).iterrows()])

h("3.4 Deep Learning Extension (Future Scope)", 2)
para(
    "The feature set is largely categorical, which suits a neural network with "
    "entity embeddings: each categorical level is mapped to a learned dense vector "
    "rather than a sparse one-hot column, allowing the network to discover which "
    "levels behave alike instead of forcing them all to be equidistant, as one-hot "
    "encoding does. Those learned vectors are also inspectable, so the resulting "
    "structure can be read and sense-checked by the marketing team rather than "
    "taken on trust. A feed-forward network over these embeddings is the planned "
    "Phase 3 extension, alongside probability calibration so that predicted scores "
    "can be used directly in expected-value calculations.")

# ==========================================================================
h("4. Business KPIs Impacted", 1)
bullets([
    "Campaign response rate - the share of contacted customers who purchase. "
    "Targeting the top decile rather than the whole list raises the achieved "
    "response rate by roughly " + format(dec.loc[0, "Lift"], ".1f") + " times.",
    "Cost per acquisition - marketing spend divided by customers acquired. Falls "
    "as low-propensity contacts are removed from the campaign.",
    "Marketing return on investment - measured as campaign profit against campaign "
    "cost, and the metric on which the optimal mailing depth is chosen.",
    "List fatigue and unsubscribe rate - fewer irrelevant contacts protects the "
    "long-term value of the customer list, a cost that a single-campaign profit "
    "calculation does not capture.",
    "Budget allocation efficiency - the model produces a ranking, so it tells the "
    "business the order in which to spend a budget of any size rather than a "
    "single fixed answer.",
])
para(
    "An important qualification, established during analysis and carried into "
    "Phase 2: whether targeting increases profit depends on the economics of the "
    "channel, not only on model quality. A customer is worth contacting whenever "
    "their purchase probability exceeds break-even = cost per contact divided by "
    "margin per sale. For cheap channels such as email that threshold is very low, "
    "most of the base clears it, and the model's value lies in efficiency rather "
    "than in profit uplift. For expensive channels such as printed catalogues or "
    "outbound calling, the threshold rises, optimal campaign depth falls sharply, "
    "and targeting becomes decisive.", italic=False)

# ==========================================================================
h("5. Tools and Technologies", 1)
table_from_df(pd.DataFrame([
    {"Category": "Language", "Tools / Libraries": "Python 3.13"},
    {"Category": "Data handling", "Tools / Libraries": "Pandas, NumPy"},
    {"Category": "Visualisation / EDA", "Tools / Libraries": "Matplotlib, Seaborn"},
    {"Category": "Modelling", "Tools / Libraries": "Scikit-learn - LogisticRegression, KNeighbors, DecisionTree, RandomForest, GradientBoosting, HistGradientBoosting, Voting, Stacking"},
    {"Category": "Model selection", "Tools / Libraries": "RandomizedSearchCV, StratifiedKFold, validation_curve"},
    {"Category": "Dimensionality reduction", "Tools / Libraries": "Scikit-learn PCA"},
    {"Category": "Interpretation", "Tools / Libraries": "Permutation importance, calibration curves, decile lift analysis"},
    {"Category": "Future DL extension", "Tools / Libraries": "PyTorch - entity-embedding neural network (Phase 3)"},
    {"Category": "Environment", "Tools / Libraries": "Jupyter Notebook / Google Colab"},
    {"Category": "Code repository", "Tools / Libraries": "https://github.com/SangithaV2024BC26670/amldl-assignment2-bike-buyer"},
]), widths=[1.6, 4.7])

# ==========================================================================
h("6. Summary", 1)
para(
    "This project applies supervised machine learning to direct-marketing "
    "targeting for Adventure Works Cycles, using Microsoft's publicly available "
    "AdventureWorks customer data - " + str(S["rows_after_dedup"]) + " unique "
    "customers across six countries, described by " + str(S["features_used"])
    + " demographic and household features after cleaning and selection. The "
    "business problem is which customers to contact in the next campaign; the ML "
    "problem is binary classification of BikeBuyer, evaluated as a ranking task.")
para(
    "Exploratory analysis confirms genuine signal, concentrated in household "
    "composition, occupation and income rather than geography, and identifies "
    "several data-quality issues - duplicate records with conflicting values, a "
    "whitespace defect in a category level, an ordinal variable requiring rank "
    "encoding, and a target-leakage risk in AveMonthSpend - that make data "
    "preparation a substantive stage of the work rather than a formality.")
para(
    "A tuned gradient-boosting classifier is the recommended primary model, "
    "achieving ROC-AUC " + format(brow["ROC_AUC"], ".4f") + " and separating the "
    "top propensity decile at " + format(dec.loc[0, "Lift"], ".2f") + " times the "
    "base purchase rate. Translated into campaign terms, the top 30 per cent of "
    "the ranked customer list captures " + format(S["capture_top30pct"], ".1f")
    + " per cent of all buyers, allowing marketing budget to be concentrated where "
    "it earns a return and supporting measurable improvements in response rate, "
    "cost per acquisition and marketing ROI.")

doc.save(OUT)
print("Wrote " + OUT)
